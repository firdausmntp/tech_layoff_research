"""
Script konversi Markdown → DOCX untuk draft paper Tech Layoffs.
Jalankan: python convert_to_docx.py

Dependensi: pip install python-docx
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MD_FILE = os.path.join(SCRIPT_DIR, "draft_paper_tech_layoffs.md")
DOCX_FILE = os.path.join(SCRIPT_DIR, "draft_paper_tech_layoffs.docx")
IMAGES_DIR = os.path.join(SCRIPT_DIR, "images")


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading_elm.append(shading)


def add_styled_paragraph(doc, text, style_name, bold=False, italic=False,
                         font_size=None, alignment=None, space_after=None,
                         font_name='Times New Roman', color=None):
    """Helper to add a paragraph with specific formatting."""
    p = doc.add_paragraph(style=style_name)
    
    # Process inline formatting (bold with ** and italic with *)
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        else:
            run = p.add_run(part)
    
    # Apply formatting to all runs
    for run in p.runs:
        if font_name and run.font.name != 'Consolas':
            run.font.name = font_name
        if font_size:
            run.font.size = Pt(font_size)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if color:
            run.font.color.rgb = color
    
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    
    return p


def parse_table(lines, start_idx):
    """Parse markdown table starting at start_idx. Returns (rows, end_idx)."""
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        row = lines[i].strip()
        # Skip separator rows (|---|---|)
        if re.match(r'^\|[\s\-:|]+\|$', row):
            i += 1
            continue
        cells = [c.strip() for c in row.split('|')[1:-1]]
        rows.append(cells)
        i += 1
    return rows, i


def add_table_to_doc(doc, rows):
    """Add a formatted table to the document."""
    if not rows:
        return
    
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = ''
                p = cell.paragraphs[0]
                
                # Process inline formatting
                parts = re.split(r'(\*\*.*?\*\*|`.*?`)', cell_text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Consolas'
                        run.font.size = Pt(8)
                    else:
                        run = p.add_run(part)
                
                for run in p.runs:
                    if run.font.name != 'Consolas':
                        run.font.name = 'Times New Roman'
                    if not run.font.size:
                        run.font.size = Pt(10)
                
                # Header row styling
                if i == 0:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    set_cell_shading(cell, '2B579A')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_code_block(doc, code_text, language=''):
    """Add a formatted code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)


def add_image_placeholder(doc, caption):
    """Add image or placeholder."""
    # Try to find the image file
    img_match = re.search(r'\((images/[^)]+)\)', caption)
    if img_match:
        img_path = os.path.join(SCRIPT_DIR, img_match.group(1))
        cap_text = re.search(r'!\[(.*?)\]', caption)
        cap_str = cap_text.group(1) if cap_text else caption
        
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Placeholder box
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'[ {cap_str} ]')
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.italic = True
        
        # Add caption
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(cap_str)
        cap_run.font.name = 'Times New Roman'
        cap_run.font.size = Pt(10)
        cap_run.italic = True


def convert_md_to_docx():
    """Main conversion function."""
    doc = Document()
    
    # ─── Page Setup ───
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.54)
    
    # ─── Default Style ───
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    # ─── Read markdown ───
    with open(MD_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_code_block = False
    code_buffer = []
    code_lang = ''
    list_counter = 0
    
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()
        
        # ─── Code Blocks ───
        if stripped.startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_buffer), code_lang)
                code_buffer = []
                in_code_block = False
                code_lang = ''
            else:
                in_code_block = True
                code_lang = stripped[3:].strip()
            i += 1
            continue
        
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
        
        # ─── Skip horizontal rules ───
        if stripped == '---':
            i += 1
            continue
        
        # ─── Empty lines ───
        if not stripped:
            i += 1
            list_counter = 0
            continue
        
        # ─── Images ───
        if stripped.startswith('!['):
            add_image_placeholder(doc, stripped)
            i += 1
            continue
        
        # ─── Title (# ) ───
        if stripped.startswith('# ') and not stripped.startswith('## '):
            title_text = stripped[2:].strip()
            # Title as centered, bold, larger
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(24)
            run = p.add_run(title_text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            i += 1
            continue
        
        # ─── Heading 2 (## ) — Bab ───
        if stripped.startswith('## ') and not stripped.startswith('### '):
            heading_text = stripped[3:].strip()
            h = doc.add_heading(heading_text, level=1)
            for run in h.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue
        
        # ─── Heading 3 (### ) — Sub-bab ───
        if stripped.startswith('### ') and not stripped.startswith('#### '):
            heading_text = stripped[4:].strip()
            h = doc.add_heading(heading_text, level=2)
            for run in h.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue
        
        # ─── Tables ───
        if stripped.startswith('|'):
            rows, end_idx = parse_table(lines, i)
            add_table_to_doc(doc, rows)
            i = end_idx
            continue
        
        # ─── Numbered lists ───
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match:
            list_counter += 1
            text = num_match.group(2)
            add_styled_paragraph(doc, text, 'List Number', font_size=12)
            i += 1
            continue
        
        # ─── Bullet lists (- or *) ───
        bullet_match = re.match(r'^[-*]\s+(.*)', stripped)
        if bullet_match:
            text = bullet_match.group(1)
            add_styled_paragraph(doc, text, 'List Bullet', font_size=12)
            i += 1
            continue
        
        # ─── Bold sub-headings (**text:**) on its own line ───
        if stripped.startswith('**') and stripped.endswith('**'):
            text = stripped[2:-2]
            add_styled_paragraph(doc, text, 'Normal', bold=True, font_size=12,
                                 space_after=4)
            i += 1
            continue
        
        # ─── Regular paragraph ───
        # Handle multi-line paragraphs (continuation lines)
        para_text = stripped
        while (i + 1 < len(lines) and
               lines[i + 1].strip() and
               not lines[i + 1].strip().startswith(('#', '|', '```', '![', '---', '-', '*')) and
               not re.match(r'^\d+\.', lines[i + 1].strip())):
            i += 1
            para_text += ' ' + lines[i].strip()
        
        add_styled_paragraph(doc, para_text, 'Normal', font_size=12)
        i += 1
    
    # ─── Save ───
    doc.save(DOCX_FILE)
    print(f"[OK] DOCX berhasil dibuat: {DOCX_FILE}")


if __name__ == '__main__':
    convert_md_to_docx()
